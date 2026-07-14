from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT_PDF = ROOT / "output" / "pdf"
OUT_DOCX = ROOT / "output" / "docx"
TMP = ROOT / "tmp" / "proposal-contract"
OUT_PDF.mkdir(parents=True, exist_ok=True)
OUT_DOCX.mkdir(parents=True, exist_ok=True)
TMP.mkdir(parents=True, exist_ok=True)

PROPOSAL_PDF = OUT_PDF / "propuesta-diseno-logo-web-panama-collectors.pdf"
CONTRACT_PDF = OUT_PDF / "contrato-diseno-logo-web-hosting-panama-collectors.pdf"
CONTRACT_DOCX = OUT_DOCX / "contrato-diseno-logo-web-hosting-panama-collectors.docx"

FONT = ROOT / "assets" / "fonts" / "PlusJakartaSans-Variable.ttf"
LOGO_WEBP = ROOT / "assets" / "panama-collectors-logo-cropped.webp"
WORDMARK_WEBP = ROOT / "assets" / "panama-collectors-wordmark.webp"
ISOTYPE_SVG = ROOT / "assets" / "panama-collectors-isotipo.svg"

ORANGE = HexColor("#f47721")
INK = HexColor("#151515")
MUTED = HexColor("#62656d")
SOFT = HexColor("#f5f5f2")
LINE = HexColor("#deded8")
WHITE = HexColor("#ffffff")
BLACK = HexColor("#0f0f10")

PROJECT_TOTAL = 320
LOGO_PRICE = 100
WEB_PRICE = 220
DEPOSIT = PROJECT_TOTAL * 0.40
BALANCE = PROJECT_TOTAL - DEPOSIT
HOSTING_MONTHLY = 20
EXTRA_ROUND = 50


def prepare_assets() -> dict[str, Path]:
    assets: dict[str, Path] = {}
    for name, src in [("logo", LOGO_WEBP), ("wordmark", WORDMARK_WEBP)]:
        img = Image.open(src).convert("RGBA")
        bbox = img.getbbox()
        if bbox:
            img = img.crop(bbox)
        out = TMP / f"{name}.png"
        img.save(out)
        assets[name] = out
    return assets


def register_fonts() -> str:
    pdfmetrics.registerFont(TTFont("PlusJakarta", str(FONT)))
    return "PlusJakarta"


def draw_wrapped(c: canvas.Canvas, text: str, x: float, y: float, width: float, size: float, color=INK, leading=None):
    leading = leading or size * 1.35
    style = ParagraphStyle(
        "p",
        fontName="PlusJakarta",
        fontSize=size,
        leading=leading,
        textColor=color,
        alignment=TA_LEFT,
    )
    p = Paragraph(text, style)
    _, h = p.wrap(width, 1000)
    p.drawOn(c, x, y - h)
    return y - h


def cover_header(c: canvas.Canvas, title: str, subtitle: str, page: int, w: float, h: float, assets: dict[str, Path]):
    c.setFillColor(SOFT)
    c.rect(0, 0, w, h, fill=1, stroke=0)
    c.setFillColor(WHITE)
    c.roundRect(28, h - 64, w - 56, 42, 14, fill=1, stroke=0)
    c.drawImage(str(assets["wordmark"]), 46, h - 54, width=140, height=28, preserveAspectRatio=True, mask="auto")
    c.setFont("PlusJakarta", 8)
    c.setFillColor(MUTED)
    c.drawRightString(w - 48, h - 43, f"Panama Collectors | {page:02d}")
    c.setFillColor(ORANGE)
    c.roundRect(42, h - 104, 118, 22, 11, fill=1, stroke=0)
    c.setFillColor(WHITE)
    c.setFont("PlusJakarta", 8)
    c.drawCentredString(101, h - 97, "DOCUMENTO COMERCIAL")
    c.setFillColor(INK)
    c.setFont("PlusJakarta", 22)
    c.drawString(42, h - 142, title)
    draw_wrapped(c, subtitle, 42, h - 158, w - 84, 9.5, MUTED, 13)


def card(c: canvas.Canvas, x: float, y: float, w: float, h: float, title: str, body: str, accent=ORANGE, fill=WHITE):
    c.setFillColor(fill)
    c.roundRect(x, y - h, w, h, 12, fill=1, stroke=0)
    c.setStrokeColor(LINE)
    c.roundRect(x, y - h, w, h, 12, fill=0, stroke=1)
    c.setFillColor(accent)
    c.roundRect(x + 14, y - 24, 28, 28, 8, fill=1, stroke=0)
    c.setFillColor(WHITE)
    c.setFont("PlusJakarta", 9)
    c.drawCentredString(x + 28, y - 17, "+")
    c.setFillColor(INK)
    c.setFont("PlusJakarta", 12)
    c.drawString(x + 54, y - 20, title)
    draw_wrapped(c, body, x + 18, y - 48, w - 36, 8.3, MUTED, 11.2)


def build_proposal(assets: dict[str, Path]):
    register_fonts()
    w, h = landscape(A4)
    c = canvas.Canvas(str(PROPOSAL_PDF), pagesize=(w, h))
    c.setTitle("Propuesta - Logo y sitio web Panama Collectors")
    c.setAuthor("Edwin Rivera")

    cover_header(
        c,
        "Propuesta de diseño de logo y sitio web",
        "Identidad visual asistida por IA, sistema de archivos SVG y desarrollo web bilingüe con 24 páginas totales.",
        1,
        w,
        h,
        assets,
    )
    c.drawImage(str(assets["logo"]), 486, 140, width=280, height=160, preserveAspectRatio=True, mask="auto")
    c.setFillColor(BLACK)
    c.roundRect(42, 92, 380, 170, 20, fill=1, stroke=0)
    c.setFillColor(WHITE)
    c.setFont("PlusJakarta", 30)
    c.drawString(66, 204, "B/.320.00")
    draw_wrapped(c, "Precio cerrado para diseño de logo, entregables SVG, página web en Next.js, 24 páginas bilingües, configuración DNS, publicación y soporte de salida.", 68, 184, 320, 9.5, HexColor("#e7e7e0"), 13)
    c.setFillColor(ORANGE)
    c.roundRect(66, 112, 142, 30, 15, fill=1, stroke=0)
    c.setFillColor(WHITE)
    c.setFont("PlusJakarta", 9)
    c.drawCentredString(137, 122, "Abono 40%: B/.128")
    c.showPage()

    cover_header(c, "Alcance creativo", "Logo diseñado con apoyo de IA y refinado para uso comercial básico.", 2, w, h, assets)
    card(c, 42, 365, 230, 122, "Sistema de logo", "Diseño de logo principal y preparación de versiones para uso digital: positivo, negativo, isotipo solo, wordmark y variantes de lectura.", ORANGE)
    card(c, 305, 365, 230, 122, "Entrega en SVG", "Se entregan archivos vectoriales SVG para uso web y escalabilidad. También se podrán preparar PNG de apoyo si se requieren para redes o documentos.", BLACK)
    card(c, 568, 365, 230, 122, "Guía de uso", "Incluye indicaciones de implementación: fondos permitidos, tamaños mínimos, separación visual, usos incorrectos y tipografía recomendada.", ORANGE)
    card(c, 42, 216, 360, 112, "Por qué el bloque de logo tiene valor B/.100", "Aunque se utilice IA como apoyo creativo, el trabajo incluye dirección visual, selección, limpieza, adaptación, preparación de versiones, exportación SVG y guía de uso. Ese bloque se valora en B/.100.00 dentro del total.", ORANGE, HexColor("#fff7f0"))
    card(c, 438, 216, 360, 112, "Tipografía", "Se documenta la familia tipográfica aprobada para el sitio y la marca. Para esta etapa se contempla Plus Jakarta Sans como base digital, con fallback seguro para navegadores.", BLACK, WHITE)
    c.showPage()

    cover_header(c, "Alcance web", "Sitio en Next.js con rutas individuales, español e inglés, WhatsApp y formulario.", 3, w, h, assets)
    data = [
        ["Bloque", "Incluye"],
        ["Páginas", "24 páginas totales: 12 en español y 12 en inglés."],
        ["Tecnología", "Implementación en Next.js, estructura responsive, SEO base, metadata y assets optimizados."],
        ["Conversión", "Botones hacia WhatsApp en navegación, páginas de servicio, CTA y footer; formulario de contacto incluido."],
        ["Servicios", "Páginas individuales para cada servicio, con estructura clara: hero, alcance, proceso, contacto por WhatsApp y formulario."],
        ["Idioma", "Estructura bilingüe para duplicar contenido ES/EN; traducción final sujeta a aprobación del cliente."],
    ]
    table = Table([[Paragraph(str(cell), ParagraphStyle("t", fontName="PlusJakarta", fontSize=8.7, leading=11, textColor=WHITE if r == 0 else INK)) for cell in row] for r, row in enumerate(data)], colWidths=[1.5 * inch, 8.6 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BLACK),
        ("BACKGROUND", (0, 1), (-1, -1), WHITE),
        ("GRID", (0, 0), (-1, -1), 0.4, LINE),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 9),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
    ]))
    table.wrapOn(c, 720, 300)
    table.drawOn(c, 42, 145)
    c.showPage()

    cover_header(c, "Configuración técnica", "Dominio existente, DNS, hosting y publicación.", 4, w, h, assets)
    card(c, 42, 360, 230, 122, "Dominio del cliente", "El cliente ya cuenta con dominio. Se incluye configuración DNS necesaria para apuntarlo al hosting o plataforma de publicación.", ORANGE)
    card(c, 305, 360, 230, 122, "DNS y SSL", "Configuración de registros DNS, conexión de dominio, certificado SSL y verificación de publicación básica.", BLACK)
    card(c, 568, 360, 230, 122, "Hosting gestionado", "Hosting mensual B/.20.00, pagadero cada día 15. Incluye mantenimiento técnico razonable y continuidad de publicación.", ORANGE)
    card(c, 42, 210, 360, 108, "No incluido", "Dominio nuevo, correo corporativo, servicios premium de terceros, fotografía profesional, campañas pagadas, CRM avanzado, textos legales especializados o cambios de alcance.", BLACK, WHITE)
    card(c, 438, 210, 360, 108, "Ajustes posteriores", "Una vez entregado el proyecto, cualquier ronda adicional de ajustes tendrá un costo de B/.50.00, salvo correcciones de errores dentro del alcance aprobado.", ORANGE, HexColor("#fff7f0"))
    c.showPage()

    cover_header(c, "Inversión y pagos", "Condiciones comerciales para iniciar y cerrar implementación.", 5, w, h, assets)
    rows = [
        ["Concepto", "Monto", "Condición"],
        ["Diseño de logo asistido por IA", "B/.100.00", "Incluye versiones SVG, positivo, negativo, isotipo, wordmark y guía de uso."],
        ["Sitio web Next.js bilingüe", "B/.220.00", "24 páginas totales, WhatsApp, formulario, DNS, SSL y publicación."],
        ["Total implementación", "B/.320.00", "Precio cerrado según alcance descrito."],
        ["Abono 40%", "B/.128.00", "Para iniciar implementación."],
        ["Saldo 60%", "B/.192.00", "Al terminar, antes de entrega final/publicación definitiva."],
        ["Hosting mensual", "B/.20.00", "Cada día 15 de cada mes."],
        ["Ronda adicional", "B/.50.00", "Aplica después de entregado el proyecto o fuera de rondas incluidas."],
    ]
    table = Table([[Paragraph(str(cell), ParagraphStyle("p", fontName="PlusJakarta", fontSize=8.2, leading=10.6, textColor=WHITE if r == 0 else INK)) for cell in row] for r, row in enumerate(rows)], colWidths=[2.6 * inch, 1.35 * inch, 6.15 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BLACK),
        ("BACKGROUND", (0, 1), (-1, -1), WHITE),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, HexColor("#fbfbf8")]),
        ("GRID", (0, 0), (-1, -1), 0.35, LINE),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 9),
        ("RIGHTPADDING", (0, 0), (-1, -1), 9),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    table.wrapOn(c, 720, 360)
    table.drawOn(c, 42, 112)
    c.showPage()

    cover_header(c, "Cronograma y aprobación", "Ruta sugerida para entregar sin confusión.", 6, w, h, assets)
    card(c, 42, 345, 175, 112, "01. Inicio", "Firma y abono B/.128.00. Recolección de contenido, dominio y accesos.", ORANGE)
    card(c, 236, 345, 175, 112, "02. Logo", "Exploración IA, selección, refinamiento y exportación de versiones.", BLACK)
    card(c, 430, 345, 175, 112, "03. Web", "Implementación Next.js, páginas ES/EN, WhatsApp, formulario y responsive.", ORANGE)
    card(c, 624, 345, 175, 112, "04. Publicación", "DNS, SSL, pruebas, saldo B/.192.00 y entrega final.", BLACK)
    draw_wrapped(c, "La implementación depende de aprobaciones, entrega de contenido y accesos. Cambios de alcance, nuevas páginas, funcionalidades no descritas o ajustes adicionales se cotizan aparte.", 42, 174, 730, 10, MUTED, 14)
    c.save()


def p_style(name, size=10.2, bold=False, color=RGBColor(35, 35, 35)):
    return name, size, bold, color


def set_font(run, size=10.2, bold=False, color=RGBColor(35, 35, 35)):
    run.font.name = "Plus Jakarta Sans"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Plus Jakarta Sans")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Plus Jakarta Sans")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_para(doc: Document, text: str, size=10.2, bold=False, color=RGBColor(35, 35, 35), align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size, bold, color)
    return p


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_font(r, 15 if level == 1 else 12.5, True, RGBColor(244, 119, 33))
    return p


def build_contract_docx(assets: dict[str, Path]):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    header = sec.header.paragraphs[0]
    header.text = ""
    hr = header.add_run("PANAMA COLLECTORS - CONTRATO DE DISEÑO, WEB Y HOSTING")
    set_font(hr, 8, True, RGBColor(100, 100, 100))
    footer = sec.footer.paragraphs[0]
    footer.text = ""
    fr = footer.add_run("Documento comercial sujeto a revisión legal antes de firma.")
    set_font(fr, 8, False, RGBColor(100, 100, 100))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(assets["logo"]), width=Inches(3.2))
    add_para(doc, "CONTRATO DE PRESTACIÓN DE SERVICIOS", 18, True, RGBColor(15, 15, 16), WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Diseño de logo asistido por IA, desarrollo web Next.js y hosting gestionado", 11, True, RGBColor(244, 119, 33), WD_ALIGN_PARAGRAPH.CENTER)

    intro = (
        "Entre los suscritos, por una parte, Edwin Rivera, cédula 8-909-1277, en adelante EL PRESTADOR; "
        "y por la otra parte, ________________________________, cédula/RUC ________________________________, "
        "en adelante EL CLIENTE; se celebra el presente contrato de prestación de servicios para Panama Collectors, "
        "sujeto a las cláusulas siguientes."
    )
    add_para(doc, intro)

    clauses = [
        ("PRIMERA. Objeto", "EL PRESTADOR diseñará una identidad visual inicial para Panama Collectors con apoyo de herramientas de inteligencia artificial, preparará sus versiones de uso y desarrollará un sitio web en Next.js con estructura bilingüe, botones hacia WhatsApp, formulario y configuración de publicación."),
        ("SEGUNDA. Alcance del logo", "El bloque de identidad tiene un valor de B/.100.00 dentro del precio total. Incluye dirección creativa con apoyo de IA, selección y refinamiento del logo, y entrega de versiones en SVG: logo principal, positivo, negativo, isotipo solo, wordmark o versión textual, variantes para fondo claro y oscuro, y archivos de apoyo que sean razonables para uso digital. También se entregará una guía breve de implementación con reglas de uso, fondos permitidos, separación visual, tamaños mínimos, usos incorrectos y tipografía recomendada."),
        ("TERCERA. Alcance web", "El sitio será desarrollado en Next.js e incluirá 24 páginas totales: 12 páginas en español y 12 páginas en inglés. Cada página de servicio deberá incluir estructura propia, contenido aprobado, llamados a WhatsApp y/o formulario. La traducción final, datos legales, nombres de servicios y textos sensibles deberán ser aprobados por EL CLIENTE antes de publicación."),
        ("CUARTA. WhatsApp y formulario", "El proyecto incluirá botones de contacto hacia WhatsApp en navegación, páginas internas, llamados a la acción y footer. También incluirá formulario de contacto. La conexión avanzada con CRM, base de datos, automatizaciones o proveedores externos no está incluida salvo acuerdo adicional."),
        ("QUINTA. Dominio, DNS y publicación", "EL CLIENTE declara que ya cuenta con dominio. EL PRESTADOR realizará la configuración necesaria de DNS, conexión de dominio, SSL y publicación en la plataforma acordada. EL CLIENTE deberá entregar accesos, permisos o colaboración técnica necesaria para completar dicha configuración."),
        ("SEXTA. Precio de implementación", f"El precio total de implementación es B/.{PROJECT_TOTAL:.2f}. Este monto se compone de B/.{LOGO_PRICE:.2f} por el bloque de logo asistido por IA y B/.{WEB_PRICE:.2f} por desarrollo web, estructura bilingüe, WhatsApp, formulario, DNS, SSL y publicación básica."),
        ("SÉPTIMA. Forma de pago", f"EL CLIENTE pagará un abono inicial del 40%, equivalente a B/.{DEPOSIT:.2f}, para iniciar el proyecto. El saldo del 60%, equivalente a B/.{BALANCE:.2f}, será pagado al terminar la implementación y antes de la entrega final, publicación definitiva, transferencia de archivos o entrega de accesos finales."),
        ("OCTAVA. Hosting mensual", f"El hosting gestionado tendrá un costo de B/.{HOSTING_MONTHLY:.2f} mensuales, pagadero cada día 15 de cada mes. Incluye alojamiento, continuidad de publicación, SSL y soporte técnico razonable relacionado con disponibilidad del sitio. No incluye rediseños, nuevas páginas, campañas, correo corporativo, dominio nuevo ni servicios premium de terceros."),
        ("NOVENA. Rondas de ajustes", f"El proyecto incluye ajustes razonables durante la implementación hasta la aprobación final. Una vez entregado el proyecto, o cuando el cambio solicitado sea una ronda adicional fuera del alcance aprobado, cada ronda adicional tendrá un costo de B/.{EXTRA_ROUND:.2f}. Cambios mayores, nuevas secciones, nuevas páginas o funcionalidades se cotizarán aparte."),
        ("DÉCIMA. Obligaciones del cliente", "EL CLIENTE entregará contenido, información de contacto, accesos de dominio, aprobaciones, logos bancarios autorizados, textos en español e inglés o aprobación de traducciones, y cualquier material necesario para completar el proyecto. Retrasos en entrega de materiales podrán mover la fecha de entrega sin constituir incumplimiento de EL PRESTADOR."),
        ("DÉCIMA PRIMERA. Propiedad y uso de entregables", "Una vez pagado el total de implementación y saldos vencidos, EL CLIENTE podrá usar los archivos del logo y el sitio para su operación comercial. EL PRESTADOR conserva herramientas, métodos, componentes reutilizables y conocimiento general no exclusivo. EL CLIENTE será responsable por el uso legal de nombres, marcas, imágenes, bancos, testimonios y textos publicados."),
        ("DÉCIMA SEGUNDA. Exclusiones", "No se incluyen asesoría legal, registro de marca, fotografía profesional, redacción legal especializada, correos corporativos, dominio nuevo, suscripciones premium, campañas publicitarias, SEO garantizado, CRM avanzado, pasarelas de pago, automatizaciones complejas ni mantenimiento de contenido recurrente."),
        ("DÉCIMA TERCERA. Aceptación", "La aprobación escrita, firma física, firma electrónica, pago del abono o confirmación por medio digital se considerará aceptación de estas condiciones comerciales."),
    ]
    for title, body in clauses:
        add_heading(doc, title, 2)
        add_para(doc, body)

    add_heading(doc, "Firmas", 1)
    add_para(doc, "Firmado en Panamá, a los ______ días del mes de __________________ de 2026.")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]
    left.text = ""
    right.text = ""
    for cell, text in [
        (left, "EL PRESTADOR\n\n\nFirma: ______________________________\nEdwin Rivera\nCédula: 8-909-1277\nFecha: ______________________________"),
        (right, "EL CLIENTE\n\n\nFirma: ______________________________\nNombre/Razón social: ______________________________\nCédula/RUC: ______________________________\nFecha: ______________________________"),
    ]:
        for i, line in enumerate(text.split("\n")):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(line)
            set_font(r, 9.2, i == 0, RGBColor(35, 35, 35))
    doc.save(CONTRACT_DOCX)


def build_contract_pdf(assets: dict[str, Path]):
    register_fonts()
    pdf = SimpleDocTemplate(str(CONTRACT_PDF), pagesize=A4, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=44)
    styles = {
        "title": ParagraphStyle("title", fontName="PlusJakarta", fontSize=18, leading=22, textColor=INK, alignment=TA_CENTER, spaceAfter=8),
        "sub": ParagraphStyle("sub", fontName="PlusJakarta", fontSize=10.5, leading=13, textColor=ORANGE, alignment=TA_CENTER, spaceAfter=14),
        "h": ParagraphStyle("h", fontName="PlusJakarta", fontSize=11.5, leading=14, textColor=ORANGE, spaceBefore=8, spaceAfter=4),
        "b": ParagraphStyle("b", fontName="PlusJakarta", fontSize=8.7, leading=12, textColor=INK, spaceAfter=6),
        "small": ParagraphStyle("small", fontName="PlusJakarta", fontSize=7.7, leading=10, textColor=MUTED, alignment=TA_CENTER),
    }
    story = []
    story.append(Spacer(1, 8))
    story.append(Paragraph("CONTRATO DE PRESTACIÓN DE SERVICIOS", styles["title"]))
    story.append(Paragraph("Diseño de logo asistido por IA, desarrollo web Next.js y hosting gestionado", styles["sub"]))
    text_blocks = [
        ("Comparecientes", "Entre los suscritos, por una parte, Edwin Rivera, cédula 8-909-1277, en adelante EL PRESTADOR; y por la otra parte, ________________________________, cédula/RUC ________________________________, en adelante EL CLIENTE; se celebra el presente contrato para Panama Collectors."),
        ("Objeto", "EL PRESTADOR diseñará una identidad visual inicial con apoyo de inteligencia artificial, preparará versiones de uso y desarrollará un sitio web en Next.js con estructura bilingüe, botones hacia WhatsApp, formulario y publicación."),
        ("Logo y entregables", "El bloque de identidad tiene valor de B/.100.00. Incluye logo principal, versiones positivo y negativo, isotipo solo, wordmark, variantes para fondos claros/oscuros, entrega en SVG y guía breve de implementación con tipografía, tamaños mínimos y usos recomendados."),
        ("Sitio web", "El sitio incluirá 24 páginas totales: 12 en español y 12 en inglés. Cada página tendrá estructura propia, contenido aprobado, llamados a WhatsApp y formulario o ruta de contacto."),
        ("DNS y dominio", "EL CLIENTE ya cuenta con dominio. EL PRESTADOR incluirá configuración DNS, conexión de dominio, SSL y publicación en plataforma acordada, sujeto a entrega de accesos y permisos."),
        ("Precio y pagos", f"Precio total: B/.{PROJECT_TOTAL:.2f}. Abono 40%: B/.{DEPOSIT:.2f} para iniciar. Saldo 60%: B/.{BALANCE:.2f} al terminar, antes de entrega final o publicación definitiva."),
        ("Hosting", f"Hosting gestionado: B/.{HOSTING_MONTHLY:.2f} mensuales, pagadero cada día 15 de cada mes. Incluye alojamiento, SSL, continuidad de publicación y soporte técnico razonable."),
        ("Ajustes adicionales", f"Una vez entregado el proyecto, cualquier ronda adicional de ajustes tendrá costo de B/.{EXTRA_ROUND:.2f}. Cambios mayores o nuevas funcionalidades se cotizan aparte."),
        ("Exclusiones", "No incluye registro de marca, asesoría legal, dominio nuevo, correo corporativo, fotografía profesional, campañas, CRM avanzado, pasarelas, automatizaciones complejas ni servicios premium de terceros."),
        ("Aceptación", "Firma, pago de abono o confirmación escrita por medio digital constituye aceptación del alcance, precio y condiciones."),
    ]
    for title, body in text_blocks:
        story.append(Paragraph(title, styles["h"]))
        story.append(Paragraph(body, styles["b"]))
    sig = Table([
        [Paragraph("<b>EL PRESTADOR</b><br/><br/><br/>Firma: ______________________________<br/>Edwin Rivera<br/>Cédula: 8-909-1277<br/>Fecha: ______________________________", styles["b"]),
         Paragraph("<b>EL CLIENTE</b><br/><br/><br/>Firma: ______________________________<br/>Nombre/Razón social: ______________________________<br/>Cédula/RUC: ______________________________<br/>Fecha: ______________________________", styles["b"])]
    ], colWidths=[3.25 * inch, 3.25 * inch], rowHeights=[1.6 * inch])
    sig.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, LINE),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 10),
    ]))
    story.extend([Spacer(1, 8), sig, Spacer(1, 8), Paragraph("Documento comercial sujeto a revisión legal antes de firma.", styles["small"])])
    pdf.build(story)


def main():
    assets = prepare_assets()
    build_proposal(assets)
    build_contract_docx(assets)
    build_contract_pdf(assets)
    print(PROPOSAL_PDF)
    print(CONTRACT_PDF)
    print(CONTRACT_DOCX)


if __name__ == "__main__":
    main()
